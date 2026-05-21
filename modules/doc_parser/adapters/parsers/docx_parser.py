"""
REQ-006 doc_parser — adapters/parsers/docx_parser.py

DocxParser
python-docx 기반 Word 문서 파서

수정 방향:
    - doc.paragraphs 먼저, doc.tables 나중 방식 제거
    - document.element.body를 순서대로 순회
    - paragraph/table 원문 순서 보존
    - table 데이터는 table 필드와 content(Markdown) 양쪽에 저장
    - DOCX는 page 개념이 고정되어 있지 않으므로 page=1 유지하되,
      source_ref.block_index로 순서를 보존
"""
from __future__ import annotations

from uuid import uuid4

import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from common_schemas.document import (
    ContentBlock,
    DocumentBlock,
    FileMeta,
    ParserMeta,
    SourceRef,
)

from doc_parser.domain.ports.parser_port import ParserPort


class DocxParser(ParserPort):
    """Word 문서 파서 구현체."""

    MIME_TYPE = (
        "application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.document"
    )

    def parse(
        self,
        file_path: str,
        file_meta: FileMeta,
    ) -> DocumentBlock:
        """DOCX 파싱 → DocumentBlock 반환."""
        try:
            doc = docx.Document(file_path)
        except Exception as e:
            raise RuntimeError(f"E0202: DOCX 파일 읽기 실패 — {e}") from e

        try:
            blocks: list[ContentBlock] = []
            block_index = 0

            # DOCX는 내부적으로 확정 page 정보가 없음.
            # 일단 1로 고정하되, 원문 순서는 block_index로 보존한다.
            page_num = 1

            # ── 본문 body를 원문 순서대로 순회 ──
            for kind, item in self._iter_body_blocks(doc):
                if kind == "paragraph":
                    text = item.text.strip()
                    if not text:
                        continue

                    block_type = self._detect_block_type(item)

                    blocks.append(
                        ContentBlock(
                            block_id=uuid4(),
                            block_type=block_type,
                            content=text,
                            page=page_num,
                            source_ref=SourceRef(
                                page=page_num,
                                block_index=block_index,
                            ),
                        )
                    )
                    block_index += 1

                elif kind == "table":
                    table_data = self._extract_table(item)

                    if not self._has_table_text(table_data):
                        # 빈 표도 디버깅 대상으로 남기고 싶다면 아래 continue 대신
                        # EMPTY_TABLE block을 만들 수 있음.
                        continue

                    markdown_table = self._table_to_markdown(table_data)

                    blocks.append(
                        ContentBlock(
                            block_id=uuid4(),
                            block_type="table",
                            # 결과 TXT/MD에서 바로 보이도록 content에도 넣는다.
                            content=markdown_table,
                            page=page_num,
                            table=table_data,
                            source_ref=SourceRef(
                                page=page_num,
                                block_index=block_index,
                            ),
                        )
                    )
                    block_index += 1

            # ── 헤더/푸터 보조 추출 ──
            # 업무문서에서 문서번호/보안등급/페이지표시가 헤더·푸터에 있을 수 있음.
            for hf_block in self._extract_headers_footers(doc, start_index=block_index):
                blocks.append(hf_block)
                block_index += 1

            if not blocks:
                raise RuntimeError("E0203: DOCX에서 추출된 텍스트/표가 없습니다.")

            return DocumentBlock(
                document_id=uuid4(),
                file_meta=file_meta,
                parser=ParserMeta(
                    parser_name="DocxParser",
                    parser_version="1.1.0",
                ),
                blocks=blocks,
            )

        except Exception as e:
            raise RuntimeError(f"E0203: DOCX 텍스트 추출 실패 — {e}") from e

    def supports(self, mime_type: str) -> bool:
        return mime_type == self.MIME_TYPE

    # ──────────────────────────────────────────
    # Body 순회
    # ──────────────────────────────────────────

    def _iter_body_blocks(self, doc: docx.Document):
        """DOCX body의 paragraph/table을 원문 순서대로 yield."""
        body = doc.element.body

        for child in body.iterchildren():
            if isinstance(child, CT_P):
                yield "paragraph", Paragraph(child, doc)
            elif isinstance(child, CT_Tbl):
                yield "table", Table(child, doc)

    # ──────────────────────────────────────────
    # Paragraph
    # ──────────────────────────────────────────

    def _detect_block_type(self, para: Paragraph) -> str:
        """단락 스타일 기반 block_type 판단."""
        style_name = para.style.name if para.style else ""

        if style_name.startswith("Heading"):
            return "heading"

        return "text"

    # ──────────────────────────────────────────
    # Table
    # ──────────────────────────────────────────

    def _extract_table(
        self,
        table: Table,
    ) -> list[list[str]]:
        """표 데이터 추출.

        cell.text만 단순 사용하면 내부 문단 줄바꿈이 뭉개질 수 있어
        cell.paragraphs를 순회해 텍스트를 보존한다.
        """
        rows: list[list[str]] = []

        for row in table.rows:
            cells: list[str] = []

            for cell in row.cells:
                paragraphs = [
                    p.text.strip()
                    for p in cell.paragraphs
                    if p.text and p.text.strip()
                ]

                text = "\n".join(paragraphs).strip()
                cells.append(text)

            rows.append(cells)

        return rows

    def _has_table_text(self, table_data: list[list[str]]) -> bool:
        """표 안에 실제 텍스트가 하나라도 있는지 확인."""
        return any(
            cell.strip()
            for row in table_data
            for cell in row
            if isinstance(cell, str)
        )

    def _table_to_markdown(self, table_data: list[list[str]]) -> str:
        """표 데이터를 결과 확인용 Markdown 문자열로 변환."""
        if not table_data:
            return ""

        # 행마다 컬럼 수가 다를 수 있으므로 최대 컬럼 수에 맞춤
        max_cols = max((len(row) for row in table_data), default=0)
        if max_cols == 0:
            return ""

        normalized = []
        for row in table_data:
            padded = list(row) + [""] * (max_cols - len(row))
            normalized.append([self._clean_cell_for_markdown(cell) for cell in padded])

        header = normalized[0]
        body = normalized[1:]

        lines = []
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * max_cols) + " |")

        for row in body:
            lines.append("| " + " | ".join(row) + " |")

        return "\n".join(lines)

    def _clean_cell_for_markdown(self, value: str) -> str:
        """Markdown 표 깨짐 방지용 셀 정리."""
        if value is None:
            return ""

        text = str(value)
        text = text.replace("|", "\\|")
        text = text.replace("\r\n", "<br>")
        text = text.replace("\n", "<br>")
        return text.strip()

    # ──────────────────────────────────────────
    # Header / Footer
    # ──────────────────────────────────────────

    def _extract_headers_footers(
        self,
        doc: docx.Document,
        start_index: int = 0,
    ) -> list[ContentBlock]:
        """섹션별 헤더/푸터 텍스트 추출."""
        blocks: list[ContentBlock] = []
        block_index = start_index
        page_num = 1

        for section_index, section in enumerate(doc.sections):
            header_text = self._collect_paragraph_text(section.header.paragraphs)
            footer_text = self._collect_paragraph_text(section.footer.paragraphs)

            if header_text:
                blocks.append(
                    ContentBlock(
                        block_id=uuid4(),
                        block_type="header",
                        content=header_text,
                        page=page_num,
                        source_ref=SourceRef(
                            page=page_num,
                            block_index=block_index,
                        ),
                    )
                )
                block_index += 1

            if footer_text:
                blocks.append(
                    ContentBlock(
                        block_id=uuid4(),
                        block_type="footer",
                        content=footer_text,
                        page=page_num,
                        source_ref=SourceRef(
                            page=page_num,
                            block_index=block_index,
                        ),
                    )
                )
                block_index += 1

        return blocks

    def _collect_paragraph_text(self, paragraphs) -> str:
        texts = [
            p.text.strip()
            for p in paragraphs
            if p.text and p.text.strip()
        ]
        return "\n".join(texts).strip()
