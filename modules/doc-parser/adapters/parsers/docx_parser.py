"""
REQ-006 doc-parser — adapters/parsers/docx_parser.py

DocxParser
python-docx 기반 Word 문서 파서

처리 항목:
    - 제목(heading) / 본문(text) 단락 구분
    - 표(table) 추출
    - 이미지 내 텍스트 제외 (MVP)
"""
from __future__ import annotations

from uuid import uuid4

import docx
from docx.oxml.ns import qn

from common_schemas.document import (
    ContentBlock,
    DocumentBlock,
    FileMeta,
    ParserMeta,
    SourceRef,
)

from doc_parser.domain.ports.parser_port import ParserPort


class DocxParser(ParserPort):
    """Word 문서 파서 구현체.

    지원 MIME 타입:
        application/vnd.openxmlformats-officedocument.wordprocessingml.document

    Raises:
        RuntimeError: 파일 손상 또는 읽기 실패 E0202
        RuntimeError: 텍스트 추출 실패 E0203
    """

    MIME_TYPE = (
        "application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.document"
    )

    # ──────────────────────────────────────────
    # ParserPort 구현
    # ──────────────────────────────────────────

    def parse(
        self,
        file_path: str,
        file_meta: FileMeta,
    ) -> DocumentBlock:
        """DOCX 파싱 → DocumentBlock 반환.

        Args:
            file_path: DOCX 파일 경로
            file_meta: 파일 메타데이터

        Returns:
            DocumentBlock: 파싱된 문서 블록

        Raises:
            RuntimeError: 파일 손상 (E0202), 텍스트 추출 실패 (E0203)
        """
        try:
            doc = docx.Document(file_path)
        except Exception as e:
            raise RuntimeError(f"E0202: DOCX 파일 읽기 실패 — {e}") from e

        try:
            blocks: list[ContentBlock] = []
            block_index = 0
            page_num = 1  # DOCX 는 페이지 정보 없음 → 1로 고정

            # ── 단락 파싱 ──
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                block_type = self._detect_block_type(para)
                blocks.append(
                    ContentBlock(
                        block_id=uuid4(),
                        block_type=block_type,
                        content=text,
                        page=page_num,
                        source_ref=SourceRef(
                            page=page_num,
                            block_index=block_index,
                        ),
                    )
                )
                block_index += 1

            # ── 표 파싱 ──
            for table in doc.tables:
                table_data = self._extract_table(table)
                if not table_data:
                    continue

                blocks.append(
                    ContentBlock(
                        block_id=uuid4(),
                        block_type="table",
                        content=None,
                        page=page_num,
                        table=table_data,
                        source_ref=SourceRef(
                            page=page_num,
                            block_index=block_index,
                        ),
                    )
                )
                block_index += 1

            return DocumentBlock(
                document_id=uuid4(),
                file_meta=file_meta,
                parser=ParserMeta(
                    parser_name="DocxParser",
                    parser_version="1.0.0",
                ),
                blocks=blocks,
            )

        except Exception as e:
            raise RuntimeError(f"E0203: DOCX 텍스트 추출 실패 — {e}") from e

    def supports(self, mime_type: str) -> bool:
        return mime_type == self.MIME_TYPE

    # ──────────────────────────────────────────
    # Private
    # ──────────────────────────────────────────

    def _detect_block_type(self, para: docx.text.paragraph.Paragraph) -> str:
        """단락 스타일 기반 block_type 판단.

        python-docx 스타일명:
            'Heading 1' ~ 'Heading 9' → heading
            그 외 → text
        """
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            return "heading"
        return "text"

    def _extract_table(
        self,
        table: docx.table.Table,
    ) -> list[list[str]]:
        """표 데이터 추출.

        Args:
            table: python-docx Table 객체

        Returns:
            list[list[str]]: 행 × 열 텍스트 데이터
        """
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        return rows